import pandas as pd
from docx import Document

# Function to generate emails from a CSV file
def generate_emails_from_csv(file_path, email_template, sender_name, program_name):
    # Load the CSV file
    data = pd.read_csv(file_path)
    
    # Extract relevant columns: "School Name" and "Contact Email Address"
    filtered_data = data[["School Name", "Contact Email Address"]].dropna()
    grouped_emails = filtered_data.groupby("School Name")["Contact Email Address"].apply(list)

    emails_to_send = []

    for school, email_list in grouped_emails.items():
        school_short_name = school.replace("Middle School", "").strip()
        
        # Replace placeholders in the email template
        email_body = email_template["body"]
        email_body = email_body.replace("{school name}", school)
        email_body = email_body.replace("[insert FULL name]", sender_name)
        email_body = email_body.replace("[insert beginning of school name]", school_short_name)
        email_body = email_body.replace("[insert program name]", program_name)
        
        email_subject = email_template["subject"].replace("{school_name}", school)
        email_sender = sender_name

        # Build the email dictionary
        email = {
            "to": ", ".join(email_list),  # Join all email addresses
            "subject": email_subject,
            "body": email_body,
            "from": email_sender
        }
        emails_to_send.append(email)

    return emails_to_send

# Function to save emails to a Google Docs-compatible file
def save_emails_to_docx(emails, output_file):
    # Create a new Word document
    document = Document()

    for email in emails:
        document.add_paragraph(f"To: {email['to']}")
        document.add_paragraph(f"Subject: {email['subject']}")
        document.add_paragraph(f"{email['body']}")

    # Save the document
    document.save(output_file)

# Define the email template
email_template = {
    "sender": "admin@stempathize.org",
    "subject": "Please Share For {school_name} Students: Build a Mind-Blowing Project! | STEMpathize",
    "body": "Hello {school name} Team,\n\nI’d like to follow up on the previous email regarding STEMpathize’s [insert program name] for [insert beginning of school name] students. I would love to further discuss our programs or how students can get involved with STEMpathize.\n\nThank you for your time, and I look forward to hearing back.\n\nBest regards,\n[insert FULL name]\nOutreach Team | STEMpathize"
}

# Parameters
file_path = 'input/STEM Teacher Research (MIDDLE).csv'  # Path to the uploaded CSV file
sender_name = "Trisha Thyagarajan"
program_name = "Saturday Workshop Series"
output_file = "Generated_Emails.docx"  # Output file path

# Generate emails
emails = generate_emails_from_csv(file_path, email_template, sender_name, program_name)

# Save emails to a Word document
save_emails_to_docx(emails, output_file)

print(f"Emails have been saved to {output_file}")
